import json
import requests
import docx
import os

BASE_URL = "http://127.0.0.1:5000"

def generate_resume(filename="sample_resume.docx"):
    doc = docx.Document()
    doc.add_heading("John Doe", 0)
    
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("Phone: +91 98765 43210")
    doc.add_paragraph("LinkedIn: linkedin.com/in/johndoe")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Technology in Computer Science and Engineering\nABC College of Engineering, Graduated 2024")
    
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Software Engineer Intern | XYZ Tech Solutions (Jan 2023 - June 2023)\n"
                      "- Developed and optimized robust enterprise backend microservices in Java and Spring Boot.\n"
                      "- Implemented REST API endpoints, improving server response times by 35%.\n"
                      "- Coded secure SQL and MySQL queries to manage client databases.")
    
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("Personal Portfolio Platform (React & Node.js)\n"
                      "- Designed and implemented a responsive portfolio using HTML, CSS, JavaScript, and React.\n"
                      "- Built REST API backend using Node.js and Express.js, handling over 1000 daily requests.\n"
                      "- Managed version releases with Git and GitHub.")
    
    doc.add_paragraph("Machine Learning Predictor (Python)\n"
                      "- Developed python scripts to train machine learning models for stock prediction.\n"
                      "- Solved critical data bottlenecks, reducing preprocessing time by 25%.")

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Developer Associate\n"
                      "Google Cloud Associate Cloud Engineer")
    
    doc.save(filename)
    print(f"Generated sample resume: {filename}")
    return filename

def test_flows():
    print("=== STARTING CAREERMATCH AI BACKEND API TESTS ===")
    
    # 1. Register User
    reg_payload = {
        "fullName": "John Doe",
        "email": "john.doe@example.com",
        "phone": "+91 98765 43210",
        "college": "ABC College of Engineering",
        "branch": "Computer Science",
        "gradYear": "2024",
        "password": "SecurePassword123"
    }
    
    try:
        # Delete old DB if it contains testing account (to avoid duplicate registration errors)
        db_file = os.path.join(os.path.dirname(__file__), "database.db")
        # However, to be safe we can just use a unique email
        import uuid
        test_email = f"test_user_{uuid.uuid4().hex[:6]}@example.com"
        reg_payload["email"] = test_email
        
        print(f"\n1. Testing User Registration with email: {test_email}...")
        res = requests.post(f"{BASE_URL}/api/register", json=reg_payload)
        print(f"Status Code: {res.status_code}")
        reg_data = res.json()
        print(f"Response: {json.dumps(reg_data, indent=2)}")
        
        assert res.status_code == 200, "Registration failed"
        assert reg_data["success"] is True, "Success should be True"
        user_id = reg_data["user"]["id"]
        print(f"SUCCESS: User Registered with ID: {user_id}")
        
        # 2. Login User
        print("\n2. Testing User Login...")
        login_payload = {
            "email": test_email,
            "password": "SecurePassword123"
        }
        res = requests.post(f"{BASE_URL}/api/login", json=login_payload)
        print(f"Status Code: {res.status_code}")
        login_data = res.json()
        print(f"Response: {json.dumps(login_data, indent=2)}")
        assert res.status_code == 200, "Login failed"
        assert login_data["success"] is True, "Success should be True"
        print("SUCCESS: User Logged In successfully")
        
        # 3. Generate and Upload Resume
        resume_file = generate_resume("test_resume.docx")
        print("\n3. Testing Resume Upload and ATS Parser...")
        with open(resume_file, "rb") as f:
            files = {"file": f}
            data = {"userId": user_id}
            res = requests.post(f"{BASE_URL}/api/upload-resume", files=files, data=data)
            
        print(f"Status Code: {res.status_code}")
        upload_data = res.json()
        print(f"ATS Score parsed: {upload_data.get('atsScore')}")
        print(f"Breakdown: {json.dumps(upload_data.get('breakdown'), indent=2)}")
        print(f"Extracted Skills: {upload_data.get('skills')}")
        
        assert res.status_code == 200, "Resume upload failed"
        assert upload_data["success"] is True, "Success should be True"
        # Verify the key mapping fixes are present in breakdown
        assert "keywords" in upload_data["breakdown"], "Fix missing: 'keywords' should be in breakdown"
        assert "content" in upload_data["breakdown"], "Fix missing: 'content' should be in breakdown"
        assert "metrics" in upload_data["breakdown"], "Fix missing: 'metrics' should be in breakdown"
        print("SUCCESS: Resume Uploaded, ATS Score parsed, and key mapping verified!")
        
        # 4. Get Recommended Jobs
        print("\n4. Testing Job Recommendations...")
        res = requests.get(f"{BASE_URL}/api/recommended-jobs/{user_id}")
        print(f"Status Code: {res.status_code}")
        jobs_data = res.json()
        print(f"Found {len(jobs_data)} job recommendations.")
        if jobs_data:
            print(f"Top recommendation: {jobs_data[0]['title']} at {jobs_data[0]['company']} ({jobs_data[0]['matchPercentage']}% Match)")
        assert res.status_code == 200, "Recommended jobs failed"
        print("SUCCESS: Recommended Jobs retrieved successfully")
        
        # 5. Apply to Job
        if not jobs_data:
            print("WARNING: No jobs found to test application flow.")
            return
            
        target_job_id = jobs_data[0]["id"]
        print(f"\n5. Testing Apply Now to Job ID: {target_job_id}...")
        apply_payload = {
            "userId": user_id,
            "jobId": target_job_id
        }
        res = requests.post(f"{BASE_URL}/api/apply-job", json=apply_payload)
        print(f"Status Code: {res.status_code}")
        apply_data = res.json()
        print(f"Response: {json.dumps(apply_data, indent=2)}")
        assert res.status_code == 200, "Job application failed"
        assert apply_data["success"] is True, "Success should be True"
        print("SUCCESS: Job application recorded successfully")
        
        # 6. Verify Dashboard and Stats
        print("\n6. Testing Dashboard Stats and Applied Job list...")
        res = requests.get(f"{BASE_URL}/api/dashboard/{user_id}")
        print(f"Status Code: {res.status_code}")
        dashboard_data = res.json()
        print(f"Dashboard Stats: {json.dumps(dashboard_data['stats'], indent=2)}")
        print(f"Applied Jobs count in Dashboard: {len(dashboard_data['appliedJobs'])}")
        
        assert res.status_code == 200, "Dashboard stats failed"
        assert dashboard_data["stats"]["appliedCount"] == 1, "Applied count should be 1"
        assert len(dashboard_data["appliedJobs"]) == 1, "Applied jobs list should have 1 item"
        print("SUCCESS: Dashboard verified! Applied job correctly shows in the user dashboard stats.")
        
        # Cleanup test file
        if os.path.exists(resume_file):
            os.remove(resume_file)
            
        print("\n=== ALL FLOW TESTS PASSED SUCCESSFULLY! ===")
        
    except Exception as e:
        print(f"\n[ERROR] TEST FAILED: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_flows()
