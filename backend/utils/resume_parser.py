import os
import re
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file_path):
    """Extract all text from a PDF file using pypdf."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract all text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_contact_info(text):
    """Extract Name, Email, and Phone using standard regexes and parsing heuristics."""
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3}[-.\s]?\d{4}'

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    email = emails[0] if emails else ""
    phone = phones[0] if phones else ""

    # Heuristic for Name: First 1-3 lines that aren't emails, phones, or URLs
    name = ""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        # Skip if contains email, phone, web links or standard resume section keywords
        if "@" in line or any(c in line for c in "0123456789") or "http" in line or "github" in line or "linkedin" in line:
            continue
        # Also skip if it is a short word or too long
        words = line.split()
        if 1 <= len(words) <= 4:
            name = line
            break
            
    if not name and lines:
        name = lines[0] # Fallback to first line
        
    return name, email, phone

def parse_resume(file_path):
    """Parses a resume (PDF or DOCX), extracts text, contact info, and matches core sections."""
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")

    name, email, phone = extract_contact_info(text)
    
    # Check for presence of key sections case-insensitively
    sections = {
        "education": False,
        "experience": False,
        "projects": False,
        "certifications": False
    }
    
    # Scan text lines for headings
    text_lower = text.lower()
    
    if re.search(r'\b(education|academic background|qualification)\b', text_lower):
        sections["education"] = True
    if re.search(r'\b(experience|employment|work|history|internship|internships|professional background)\b', text_lower):
        sections["experience"] = True
    if re.search(r'\b(project|projects|academic projects|personal projects)\b', text_lower):
        sections["projects"] = True
    if re.search(r'\b(certification|certifications|certificate|certificates|achievements|awards)\b', text_lower):
        sections["certifications"] = True

    return {
        "text": text,
        "name": name,
        "email": email,
        "phone": phone,
        "sections": sections
    }
